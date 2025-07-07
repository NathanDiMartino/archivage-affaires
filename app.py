import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinterdnd2 import DND_FILES, TkinterDnD
from PIL import Image, ImageTk
from tkinter import font
import threading as th
import pandas as pd
from docx import Document
import time
import os
import re

import utils
from tkinter_custom_classes_ndm_module.tkinter_custom_classes_ndm.utils import *


# FENÊTRE DE DÉMARRAGE
class Demarrage(tk.Toplevel):
    def __init__(self, master, on_finish):
        super().__init__(master)
        self.geometry("400x100+568+322")
        self.resizable(False, False)
        self.configure(background="white")
        self.on_finish = on_finish

        self.protocol("WM_DELETE_WINDOW", self.kill)

        self.label = ttk.Label(self)
        self.label.pack(expand=tk.YES)

        self.last_position = (self.winfo_x(), self.winfo_y())
        self.bind("<Configure>", self.enregistrer_position)

        thread = th.Thread(target=self.lancer_application, daemon=True)
        thread.start()

    def enregistrer_position(self, event):
        # Enregistre dynamiquement la position de la fenêtre
        if event.widget == self:
            self.last_position = (self.winfo_x(), self.winfo_y())

    def recuperer_mandats(self):
        self.label.config(text="Récupération des mandats...")
        time.sleep(1)
        chemin = "C:\\Users\\nathan.dimartino\\Documents\\Archivage affaires\\H"
        self.liste_mandats = []

        for affaire in [d for d in os.listdir(chemin) if os.path.isdir(os.path.join(chemin, d))]:
            for mandat in os.listdir(os.path.join(chemin, affaire)):
                if not mandat.startswith("_"):
                    self.liste_mandats.append(mandat)

        time.sleep(1)
        
    def recuperer_mandats_a_archiver(self):
        self.label.config(text="Récupération des mandats prêts à être archivés...")
        time.sleep(1)
        chemin = "C:\\Users\\nathan.dimartino\\Documents\\Archivage affaires\\H"
        self.liste_mandats_a_archiver = utils.liste_mandat_a_archiver(chemin=chemin)
        nb = len(self.liste_mandats_a_archiver)
        self.label.config(text=f"{nb} mandat{'s' if nb != 1 else ''} prêt{'s' if nb != 1 else ''} à être archivé{'s' if nb != 1 else ''}.")

        time.sleep(1)
        self.label.config(text="Lancement de l'application...")

    def lancer_application(self):

        self.recuperer_mandats()
        self.recuperer_mandats_a_archiver()

        time.sleep(1)

        self.destroy()
        self.on_finish(self.liste_mandats, self.liste_mandats_a_archiver, self.last_position)

    def kill(self):
        self.destroy()
        self.master.destroy()


# APPLICATION PRINCIPALE
class Application(tk.Toplevel):
    def __init__(self, master, liste_mandats, liste_mandats_a_archiver, position=None):
        super().__init__(master)

        if position:
            x, y = position
            self.geometry(f"+{x+50}+{y+50}")  # petit décalage optionnel
        else:
            self.geometry("+100+100")

        self.title("QMT Archivage")
        self.state("zoomed")
        self.minsize(480, 360)
        self.configure(background="white")

        self.protocol("WM_DELETE_WINDOW", self.kill)

        self.racine = Racine(self, liste_mandats, liste_mandats_a_archiver)
        self.racine.pack(expand=True, fill="both")

    def kill(self):
        self.destroy()
        self.master.destroy()


# CONTENU PRINCIPAL
class Racine(ttk.Frame):
    def __init__(self, master, liste_mandats, liste_mandats_a_archiver):
        super().__init__(master)
        self.master = master
        self.liste_mandats = liste_mandats
        self.liste_mandats_a_archiver = liste_mandats_a_archiver
        self.liste_mandats_a_editer = []
        self.chemin = "C:\\Users\\nathan.dimartino\\Documents\\Archivage affaires\\H"

        self.delta1 = 3
        self.delta2 = 0
        self.barre_menu()
        self.page()

        self.statut = "sleep"

    def barre_menu(self):

        self.menu = tk.Menu(self.master)

        # Fichier
        self.menu_fichier = tk.Menu(self.menu, tearoff=0)

        # Ajouter des mandats à éditer depuis un fichier
        self.menu_fichier.add_command(
            label="Ajouter des mandats à éditer depuis un fichier", 
            command=self.charger, 
            accelerator="Ctrl+O",
        )
        self.bind_all("<Control-o>", lambda event: self.charger())

        # ---
        # Quitter
        self.menu_fichier.add_separator()
        self.menu_fichier.add_command(
            label="Quitter", 
            command=self.kill, 
            accelerator="Ctrl+Q",
        )
        self.bind_all("<Control-q>", lambda event: self.kill())

        self.menu.add_cascade(
            label="Fichier",
            menu=self.menu_fichier,
        )

        # Édition
        self.menu_edition = tk.Menu(self.menu, tearoff=0)

        # Rafraîchir la liste des mandats à archiver
        self.menu_edition.add_command(
            label="Rafraîchir la liste des mandats à archiver", 
            command=self.a_archiver, 
            accelerator="F5",
        )
        self.bind_all("<F5>", lambda event: self.a_archiver())

        # Archiver les mandats
        self.menu_edition.add_command(
            label="Archiver les mandats", 
            command=self.a_archiver, 
            accelerator="Ctrl+R",
        )
        self.bind_all("<Control-r>", lambda event: self.a_archiver())

        # Ne plus archiver les mandats
        self.menu_edition.add_command(
            label="Ne plus archiver les mandats", 
            command=self.pas_a_archiver, 
            accelerator="Ctrl+Shift+R",
        )
        self.bind_all("<Control-Shift-KeyPress-R>", lambda event: self.pas_a_archiver())

        self.menu.add_cascade(
            label="Édition",
            menu=self.menu_edition,
        )

        # Sélection
        self.menu_selection = tk.Menu(self.menu, tearoff=0)

        # Ajouter des mandats à la sélection
        self.menu_selection.add_command(
            label="Ajouter des mandats à la sélection depuis un fichier",
            command=self.ajouter_selection,
            accelerator="Ctrl+Shift+O"
        )
        self.bind_all("<Control-Shift-KeyPress-O>", lambda event: self.ajouter_selection())

        # Supprimer la sélection
        self.menu_selection.add_command(
            label="Supprimer la sélection",
            command=self.supprimer_selection,
            accelerator="Ctrl+Shift+D"
        )
        self.bind_all("<Control-Shift-KeyPress-D>", lambda event: self.supprimer_selection())

        self.menu.add_cascade(
            label="Sélection",
            menu=self.menu_selection,
        )

        # Exécution
        self.menu_execution = tk.Menu(self.menu, tearoff=0)

        # Lancer l'archivage
        self.menu_execution.add_command(
            label="Lancer l'archivage", 
            command=self.archiver, 
            accelerator="Ctrl+Entrer",
        )
        self.bind_all("<Control-Return>", lambda event: self.archiver())

        # Arrêter l'archivage
        self.menu_execution.add_command(
            label="Arrêter l'archivage", 
            command=self.stop, 
            accelerator="Ctrl+S",
        )
        self.bind_all("<Control-s>", lambda event: self.stop())

        # Arrêt d'urgence
        self.menu_execution.add_command(
            label="Arrêt d'urgence", 
            command=self.kill, 
            accelerator="Ctrl+K",
        )
        self.bind_all("<Control-k>", lambda event: self.kill())

        self.menu.add_cascade(
            label="Exécution",
            menu=self.menu_execution,
        )

        self.master.config(menu=self.menu)

    # Fonctions utiles
    def charger(self):
        chemin = filedialog.askopenfile(parent=self.master, defaultextension="*.xlsx", filetypes=[
            ("Tous les fichiers", "*.xlsx *.xlsm *.csv *.CSV *.txt *.docx"),
            ("Fichiers Excel", "*.xlsx"),
            ("Fichiers Excel", "*.xlsm"),
            ("Fichiers CSV", "*.csv"),
            ("Fichiers CSV", "*.CSV"),
            ("Fichiers texte", "*.txt"),
            ("Fichiers Word", "*.docx")
        ])

        if not chemin:
            return
        chemin = chemin.name

        extension = chemin.split(".")[-1]
        separateur = None
        en_tete = False
        colonne = 0
        
        match extension:
            case x if x in ["xlsx", "xlsm"]: 
                try:
                    df = pd.read_excel(chemin)
                    colonne = 0

                    for i in range(len(df.columns)):
                        if str(df.columns[i]) in self.liste_mandats:
                            colonne = i
                            break
                        elif str(df.iloc[0, i]) in self.liste_mandats:
                            en_tete = True
                            colonne = i
                            break

                except:
                    colonne = 0

            case x if x in ["csv", "CSV"]:
                try:
                    
                    separateurs = [",", ";"]
                    bon_sep = False
                    i = 0
                    while not bon_sep:
                        colonne = 0
                        separateur = separateurs[i]
                        df = pd.read_csv(chemin, sep=separateur)
                        
                        for i in range(len(df.columns)):
                            if str(df.columns[i]) in self.liste_mandats:
                                colonne = i
                                bon_sep = True
                                break
                            elif str(df.iloc[0, i]) in self.liste_mandats:
                                en_tete = True
                                colonne = i
                                bon_sep = True
                                break
                        i += 1
                        
                except:
                    colonne = 0
                    separateur = ";"

            case "docx":
                try:
                    doc = Document(chemin)
                    texte = ""
                    for para in doc.paragraphs:
                        texte = texte + "SAUT_DE_LIGNE" + para.text

                    separateurs = re.findall(r"\d{4}-\d{2}(.*?)(?=\d{4}-\d{2})", texte)
                    if len(separateurs) == 1 or len(set(separateurs)) == 1:
                        separateur = separateurs[0]
                
                except:
                    pass

            case "txt":
                try:
                    with open(chemin, "r", encoding="utf-8") as fichier:
                        lignes = fichier.readlines()
                    texte = ""
                    for ligne in lignes:
                        texte = texte + "SAUT_DE_LIGNE" + ligne.replace("\n", "")

                    separateurs = re.findall(r"\d{4}-\d{2}(.*?)(?=\d{4}-\d{2})", texte)
                    if len(separateurs) == 1 or len(set(separateurs)) == 1:
                        separateur = separateurs[0]
                
                except:
                    pass

        popup = PopUpCharger(self.master, racine=self, on_finish=None, largeur=450, hauteur=300, chemin=chemin, separateur=separateur, en_tete=en_tete, colonne=colonne)
        popup.wait_window()
        return popup.mandats_ajoutes

    def kill(self):
        self.master.kill()

    def a_archiver(self):
        pass

    def pas_a_archiver(self):
        pass

    def ajouter_selection(self):
        liste_mandats = self.charger()
        print(liste_mandats)

    def supprimer_selection(self):
        pass

    def archiver(self):
        self.bouton_archiver.start()

        self.bouton_archiver.set_state("disabled")
        self.bouton_stop.set_state("normal")
        self.bouton_kill.set_state("normal")

        self.statut = "working"
        print(self.liste_mandats_a_archiver)

        thread = th.Thread(
            target=utils.archivage,
            args=(self.liste_mandats_a_archiver,),
            kwargs={
                "chemin": self.chemin,
                "chemin_archive": "C:\\Users\\nathan.dimartino\\Documents\\Archivage affaires\\Z\\Affaires",
                "application": self
            },
            daemon=True
        )
        thread.start()

        self.frame_colonne_droite.grid_forget()
        self.frame_colonne_progression_archivage.grid(row=1, column=1, sticky="news")

    def stop(self):
        self.bouton_stop.start()

        self.bouton_stop.set_state("disabled")
        self.bouton_kill.set_state("disabled")

        self.statut = "sleep"

        self.rafraichir()

    def rafraichir(self):

        thread = th.Thread(target=self.maj_mandats_a_archiver, daemon=True)
        thread.start()

    def maj_mandats_a_archiver(self):
        for label in self.frame_liste_mandats_a_archiver.grid_slaves():
            label.grid_forget()

        self.label_nombre_mandats.start()
        self.label_nombre_mandats.config(text=f"Détection des mandats prêts à être archivés dans le dossier \"{self.chemin}\"...")
        self.liste_mandats_a_archiver = utils.liste_mandat_a_archiver(chemin=self.chemin)
        
        self.liste_mandats = []

        for affaire in [d for d in os.listdir(self.chemin) if os.path.isdir(os.path.join(self.chemin, d))]:
            for mandat in os.listdir(os.path.join(self.chemin, affaire)):
                if not mandat.startswith("_"):
                    self.liste_mandats.append(mandat)

        self.label_nombre_mandats.stop()
        self.label_nombre_mandats.config(text=f"{len(self.liste_mandats_a_archiver)} mandat{"s" if len(self.liste_mandats_a_archiver) else ""} prêt{"s" if len(self.liste_mandats_a_archiver) else ""} à être archivés dans le dossier \"{self.chemin}\".")

        ligne = 0
        for mandat, _ in self.liste_mandats_a_archiver:
            ttk.Label(self.frame_liste_mandats_a_archiver, text=mandat).grid(row=ligne, column=0, sticky="w", ipady=self.delta1)
            ligne += 1

    def precedent_mandat(self, event):
        match event.keysym:
            case "Up":
                sens = -1
            case "Down":
                sens = 1
            case _:
                return

        try:
            _ = self.precedentes_recherches[self.indice_recherches + sens]
            self.indice_recherches += sens
            self.barre_de_recherche.delete(0, tk.END)
            self.barre_de_recherche.insert(0, self.precedentes_recherches[self.indice_recherches])

        except IndexError:
            pass

    def afficher_message_temporaire(self, texte, couleur="red", duree=2000, wait=3000):
        top = tk.Toplevel(self)
        top.overrideredirect(True)
        top.attributes("-topmost", True)
        top.attributes("-alpha", 1.0)
        top.configure(bg=couleur)

        frame = tk.Frame(
            top,
            bg=couleur,
        )
        label = tk.Text(
            frame, 
            bg=couleur, 
            fg="white", 
            font=("Segoe UI", 9, "bold"), 
            width=50, height=min(len(texte.split("\n")), 15), 
            bd=0
        )
        label.insert('1.0', texte)
        label.config(state="disabled")
        label.pack(padx=10, pady=5)
        frame.pack()

        # ➜ Mise à jour du placement principal
        self.update_idletasks()
        top.update_idletasks()

        # ➜ Centrage par rapport à la fenêtre principale sur le bon écran
        x = self.winfo_rootx() + 20
        y = self.winfo_rooty() + 20
        top.geometry(f"+{x}+{y}")

        top.after(wait, lambda: self._fade_out(top, 1.0, duree))

    def _fade_out(self, window, alpha, duree):
        steps = 200
        delay = duree // steps

        def fade():
            nonlocal alpha
            alpha -= 1 / steps
            if alpha <= 0:
                window.destroy()
            else:
                window.attributes("-alpha", alpha)
                window.after(delay, fade)

        fade()

    def ajouter_ligne_a_editer(self, mandat):
        frame = ttk.Frame(self.frame_liste_mandats_a_editer)
        frame.grid_columnconfigure(0, weight=0)
        frame.grid_columnconfigure(1, weight=1)
        frame.grid_columnconfigure(2, weight=0)

        var_selection = tk.BooleanVar()
        checkbox = ttk.Checkbutton(frame, variable=var_selection)
        checkbox.grid(row=0, column=0, padx=5)

        label = ttk.Label(frame, text=mandat)
        label.grid(row=0, column=1, sticky="w")

        btn_supprimer = ttk.Button(frame, text="x", width=3, command=lambda: self.supprimer_ligne(frame, (mandat, var_selection)))
        btn_supprimer.grid(row=0, column=2, padx=5)

        frame.pack(fill=tk.X, side=tk.TOP, anchor=tk.N, ipady=self.delta2)

        self.liste_mandats_a_editer.append((mandat, var_selection))

    def supprimer_ligne(self, frame, mandat_tuple):
        frame.destroy()
        if mandat_tuple in self.liste_mandats_a_editer:
            self.liste_mandats_a_editer.remove(mandat_tuple)

    def ajouter_mandat_a_editer(self, mandat=None):
        if mandat is None:
            mandat = self.barre_de_recherche.get()
        mandat = mandat.strip()
        mandat = mandat.replace("\n", "")

        if mandat in self.liste_mandats and mandat not in [m for m, _ in self.liste_mandats_a_editer]:
            self.ajouter_ligne_a_editer(mandat)
            self.barre_de_recherche.delete(0, tk.END)
            self.precedentes_recherches.append(mandat)
            self.indice_recherches = 0

        if mandat in [m for m, _ in self.liste_mandats_a_editer]:
            return

        candidats = [m for m in self.liste_mandats if mandat in m]

        if mandat == "":
            pass
        elif len(candidats) == 0:
            self.afficher_message_temporaire(f"Aucun mandat correspondant à \"{mandat}\".")
        else:
            self.afficher_message_temporaire(f"Aucun mandat correspondant à \"{mandat}\".\nEssayer plutôt :\n- {"\n- ".join([f"\"{m}\"" for m in candidats])}")

    def handle_file_drop(self, event):
        fichier = event.data.strip("{}")  # gère les chemins avec espaces
        self.traiter_fichier_droppe(fichier)

    def traiter_fichier_droppe(self, chemin):
        if not os.path.isfile(chemin):
            self.afficher_message_temporaire("Fichier invalide.", couleur="orange")
            return

        extension = chemin.split(".")[-1]
        separateur = None
        en_tete = False
        colonne = 0
        
        match extension:
            case x if x in ["xlsx", "xlsm"]: 
                try:
                    df = pd.read_excel(chemin)
                    colonne = 0

                    for i in range(len(df.columns)):
                        if str(df.columns[i]) in self.liste_mandats:
                            colonne = i
                            break
                        elif str(df.iloc[0, i]) in self.liste_mandats:
                            en_tete = True
                            colonne = i
                            break

                except:
                    colonne = 0

            case x if x in ["csv", "CSV"]:
                try:
                    
                    separateurs = [",", ";"]
                    bon_sep = False
                    i = 0
                    while not bon_sep:
                        colonne = 0
                        separateur = separateurs[i]
                        df = pd.read_csv(chemin, sep=separateur)
                        
                        for i in range(len(df.columns)):
                            if str(df.columns[i]) in self.liste_mandats:
                                colonne = i
                                bon_sep = True
                                break
                            elif str(df.iloc[0, i]) in self.liste_mandats:
                                en_tete = True
                                colonne = i
                                bon_sep = True
                                break
                        i += 1
                        
                except:
                    colonne = 0
                    separateur = ";"

            case "docx":
                try:
                    doc = Document(chemin)
                    texte = ""
                    for para in doc.paragraphs:
                        texte = texte + "SAUT_DE_LIGNE" + para.text

                    separateurs = re.findall(r"\d{4}-\d{2}(.*?)(?=\d{4}-\d{2})", texte)
                    if len(separateurs) == 1 or len(set(separateurs)) == 1:
                        separateur = separateurs[0]
                
                except:
                    pass

            case "txt":
                try:
                    with open(chemin, "r", encoding="utf-8") as fichier:
                        lignes = fichier.readlines()
                    texte = ""
                    for ligne in lignes:
                        texte = texte + "SAUT_DE_LIGNE" + ligne.replace("\n", "")

                    separateurs = re.findall(r"\d{4}-\d{2}(.*?)(?=\d{4}-\d{2})", texte)
                    if len(separateurs) == 1 or len(set(separateurs)) == 1:
                        separateur = separateurs[0]
                
                except:
                    pass

        popup = PopUpCharger(self.master, racine=self, on_finish=None, largeur=450, hauteur=300, chemin=chemin, separateur=separateur, en_tete=en_tete, colonne=colonne)
        popup.wait_window()
        return popup.mandats_ajoutes
    
    def changer_chemin(self):
        chemin = filedialog.askdirectory(parent=self.master, initialdir=self.chemin)

        if not chemin and chemin == "":
            return
        self.chemin = chemin
        self.rafraichir()

    def page(self):
        self.grid_rowconfigure(0, weight=0)
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=0)
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        self.frame_en_tete = self.en_tete()
        self.frame_en_tete.grid(row=0, column=0, columnspan=2, sticky="ew")
        self.frame_colonne_gauche= self.colonne_gauche()
        self.frame_colonne_gauche.grid(row=1, column=0, sticky="news")
        self.frame_colonne_droite= self.colonne_droite()
        self.frame_colonne_droite.grid(row=1, column=1, sticky="news")
        self.frame_colonne_progression_archivage = self.colonne_progression_archivage()
        self.frame_pied_de_page= self.pied_de_page()
        self.frame_pied_de_page.grid(row=2, column=0, columnspan=2, sticky="ew")

    #     self.after_idle(self.synchroniser_hauteurs)

    # def synchroniser_hauteurs(self):
    #     self.update_idletasks()
    #     h1 = self.label_nombre_mandats.winfo_height()
    #     h2 = self.frame_barre_de_recherche.winfo_height()
    #     h3 = self.barre_de_progression.winfo_height()
    #     self.h = max(h1, h2, h3)
    #     self.delta1 = (self.h - h1) // 2
    #     self.delta2 = (self.h - h2) // 2
    #     print(h1, h2, self.h)

    def en_tete(self):
        frame = ttk.Frame(self)
        label_en_tete = ttk.Label(frame, text="QMT Archivage", font=("Helvetica", 16))
        label_en_tete.pack(padx=10, pady=10)
        return frame

    def colonne_gauche(self):
        frame = ttk.Frame(self)
        frame.pack_propagate(False)
        label_colonne_gauche = ttk.Label(frame, text="Mandats à  archiver")
        label_colonne_gauche.pack(padx=10, pady=10)

        self.frame_nombre_mandats = ttk.Frame(frame)
        self.frame_nombre_mandats.grid_columnconfigure(0, weight=1)
        self.frame_nombre_mandats.grid_columnconfigure(1, weight=0)
        self.frame_nombre_mandats.grid_columnconfigure(2, weight=0)
        self.frame_nombre_mandats.pack(fill=tk.X, padx=10, pady=10)

        self.label_nombre_mandats = ChargementButton(self.frame_nombre_mandats, text=f"{len(self.liste_mandats_a_archiver)} mandat{"s" if len(self.liste_mandats_a_archiver) else ""} prêt{"s" if len(self.liste_mandats_a_archiver) else ""} à être archivés dans le dossier \"{self.chemin if len(self.chemin) <= 100 else "..." + self.chemin[-97:]}\".", padding=10, command=lambda chemin=self.chemin: os.startfile(chemin))
        self.label_nombre_mandats.grid(row=0, column=0)
        self.bouton_chemin = ttk.Button(self.frame_nombre_mandats, width=3, text="+", command=self.changer_chemin)
        self.bouton_chemin.grid(row=0, column=1, padx=(10, 5))
        self.bouton_rafraichir = ttk.Button(self.frame_nombre_mandats, width=3, text="+", command=self.rafraichir)
        self.bouton_rafraichir.grid(row=0, column=2, padx=(5, 0))

        self.frame_liste_mandats_a_archiver = ttk.Frame(frame)
        self.frame_liste_mandats_a_archiver.pack(expand=tk.YES, fill=tk.BOTH, padx=10, pady=10)

        ligne = 0
        for mandat, _ in self.liste_mandats_a_archiver:
            ttk.Label(self.frame_liste_mandats_a_archiver, text=mandat).grid(row=ligne, column=0, sticky="w", ipady=self.delta1)
            ligne += 1

        return frame

    def colonne_droite(self):
        frame = ttk.Frame(self)
        frame.pack_propagate(False)
        label_colonne_droite = ttk.Label(frame, text="Éditer des mandats")
        label_colonne_droite.pack(padx=10, pady=10)

        self.frame_barre_de_recherche = ttk.Frame(frame)
        self.frame_barre_de_recherche.grid_columnconfigure(0, weight=1)
        self.frame_barre_de_recherche.grid_columnconfigure(1, weight=0)
        self.frame_barre_de_recherche.grid_columnconfigure(2, weight=0)
        self.frame_barre_de_recherche.pack(fill=tk.X, padx=10, pady=10)

        self.barre_de_recherche = ttk.Entry(self.frame_barre_de_recherche)
        self.barre_de_recherche.grid(row=0, column=0, sticky="ew")
        self.barre_de_recherche.bind("<Return>", lambda event: self.ajouter_mandat_a_editer())
        self.barre_de_recherche.bind("<Up>", self.precedent_mandat)
        self.barre_de_recherche.bind("<Down>", self.precedent_mandat)
        self.barre_de_recherche.bind("<Key>", lambda e: setattr(self, "indice_recherches", 0))

        self.bouton_recherche = ttk.Button(self.frame_barre_de_recherche, width=3, command=self.ajouter_mandat_a_editer, text="+")
        self.bouton_recherche.grid(row=0, column=1, padx=(10, 5))
        self.precedentes_recherches = []

        self.bouton_charger = ttk.Button(self.frame_barre_de_recherche, text="+", width=3, command=self.charger)
        self.bouton_charger.grid(row=0, column=2, padx=(5, 0))

        self.frame_liste_mandats_a_editer = ttk.Frame(frame)
        self.frame_liste_mandats_a_editer.pack_propagate(False)
        self.frame_liste_mandats_a_editer.pack(expand=tk.YES, fill=tk.BOTH, padx=10, pady=10, anchor=tk.N)
        self.liste_mandats_a_editer = []

        self.frame_liste_mandats_a_editer.drop_target_register(DND_FILES)
        self.frame_liste_mandats_a_editer.dnd_bind('<<Drop>>', self.handle_file_drop)

        return frame
    
    def colonne_progression_archivage(self):
        frame = ttk.Frame(self)
        frame.pack_propagate(False)
        label_colonne_droite = ttk.Label(frame, text="Progression de l'archivage")
        label_colonne_droite.pack(padx=10, pady=10)

        self.frame_barre_de_progression = ttk.Frame(frame)
        self.frame_barre_de_progression.grid_columnconfigure(0, weight=1)
        self.frame_barre_de_progression.grid_columnconfigure(1, weight=0)
        self.frame_barre_de_progression.grid_columnconfigure(2, weight=0)
        self.frame_barre_de_progression.pack(fill=tk.X, padx=10, pady=10, ipady=self.delta2)

        self.barre_de_progression = ttk.Progressbar(self.frame_barre_de_progression, orient=tk.HORIZONTAL)
        self.barre_de_progression.grid(row=0, column=0, sticky="ew")

        self.label_barre_de_progression = ttk.Label(self.frame_barre_de_progression, text="-/-")
        self.label_barre_de_progression.grid(row=0, column=1, padx=10, ipady=self.delta1)

        self.frame_log = tk.Text(frame)
        self.frame_log.pack(expand=tk.YES, fill=tk.BOTH, padx=10, pady=10, anchor=tk.N)

        return frame

    def pied_de_page(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)
        frame.grid_columnconfigure(2, weight=1)

        self.bouton_archiver = ChargementButton(frame, text="Archiver", command=self.archiver, state="normal" if len(self.liste_mandats_a_archiver) > 0 else "disabled", padding=10)
        self.bouton_archiver.grid(row=0, column=0, sticky="", padx=10, pady=10)
        self.bouton_stop = ChargementButton(frame, text="Stop", command=self.stop, state="disabled", padding=10)
        self.bouton_stop.grid(row=0, column=1, sticky="", padx=10, pady=10)
        self.bouton_kill = ChargementButton(frame, text="Arrêt d'urgence", command=self.kill, state="disabled", padding=10)
        self.bouton_kill.grid(row=0, column=2, sticky="", padx=10, pady=10)
        
        return frame


class PopUpCharger(PopUp):
    def __init__(self, master, racine, on_finish, largeur, hauteur, chemin, separateur, en_tete, colonne):
        self.chemin = chemin
        self.separateur = separateur
        self.var_en_tete = tk.BooleanVar(value=en_tete)
        self.colonne = colonne
        self.mandats_ajoutes = None

        self.extension = self.chemin.split(".")[-1]

        super().__init__(master, racine, on_finish, largeur, hauteur)

    def page(self):
        self.grid_rowconfigure(0, weight=0)
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=0)
        self.grid_columnconfigure(0, weight=1)

        frame_en_tete = self.en_tete()
        frame_en_tete.grid(row=0, column=0, sticky="news")
        frame_pied_de_page = self.pied_de_page()
        frame_pied_de_page.grid(row=2, column=0, sticky="news")

        if self.extension in ["xlsx", "xlsm"]:
            self.frame_corps = self.page_excel()
        elif self.extension in ["csv", "CSV"]:
            self.frame_corps = self.page_csv()
        else:
            self.frame_corps = self.page_texte()

        self.frame_corps.grid(row=1, column=0, sticky="new", padx=10, pady=10)

    def en_tete(self):
        frame = ttk.Frame(self)
        label_en_tete = ttk.Label(frame, text="Charger un fichier", font=("Helvetica", 16))
        label_en_tete.pack(padx=10, pady=10)
        return frame

    def page_excel(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=0)
        frame.grid_rowconfigure(1, weight=1)
        frame.grid_rowconfigure(2, weight=1)
        frame.grid_rowconfigure(3, weight=1)
        frame.grid_columnconfigure(0, weight=0)
        frame.grid_columnconfigure(1, weight=0)
        frame.grid_columnconfigure(2, weight=1)

        self.label_nombre_mandats = ttk.Label(frame)
        self.label_nombre_mandats.grid(row=0, column=0, columnspan=3, padx=10, pady=(0, 10))

        thread = th.Thread(target=self.detection_des_mandats, daemon=True)
        thread.start()

        label_chemin = ttk.Label(frame, text="Fichier")
        label_chemin.grid(row=1, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=1, column=1, sticky="ew", padx=10)

        bouton_chemin = tk.Button(frame, text=self.chemin.split("/")[-1], anchor="w", command=lambda chemin=self.chemin: os.startfile(chemin))
        bouton_chemin.grid(row=1, column=2, sticky="ew")

        label_colonne = ttk.Label(frame, text="Colonne")
        label_colonne.grid(row=2, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=2, column=1, sticky="ew", padx=10)

        df = pd.read_excel(self.chemin, header=0 if self.var_en_tete.get() else None)
        self.menu_colonne = MenuDeroulant(frame, taille=5, items=df.columns, text=df.columns[self.colonne], command=self.detection_des_mandats)
        self.menu_colonne.selection_index = self.colonne
        self.menu_colonne.grid(row=2, column=2, sticky="ew")

        frame_en_tete = ttk.Frame(frame)

        frame_en_tete.grid_columnconfigure(0, weight=0)
        frame_en_tete.grid_columnconfigure(1, weight=1)
        label_en_tete = ttk.Label(frame_en_tete, text="Mon fichier comporte des entêtes :")
        label_en_tete.grid(row=0, column=0, sticky="w")

        checkbox_en_tete = ttk.Checkbutton(frame_en_tete, variable=self.var_en_tete)
        checkbox_en_tete.grid(row=0, column=1, sticky="w")
        checkbox_en_tete.bind("<Button-1>",  lambda event: self.modifier_checkbox(self.var_en_tete))

        frame_en_tete.grid(row=3, column=0, columnspan=3, sticky="ew")

        return frame
    
    def page_csv(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=0)
        frame.grid_rowconfigure(1, weight=1)
        frame.grid_rowconfigure(2, weight=1)
        frame.grid_rowconfigure(3, weight=1)
        frame.grid_rowconfigure(4, weight=1)
        frame.grid_columnconfigure(0, weight=0)
        frame.grid_columnconfigure(1, weight=0)
        frame.grid_columnconfigure(2, weight=1)

        self.label_nombre_mandats = ttk.Label(frame)
        self.label_nombre_mandats.grid(row=0, column=0, columnspan=3, padx=10, pady=(0, 10))

        thread = th.Thread(target=self.detection_des_mandats, daemon=True)
        thread.start()

        label_chemin = ttk.Label(frame, text="Fichier")
        label_chemin.grid(row=1, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=1, column=1, sticky="ew", padx=10)

        bouton_chemin = tk.Button(frame, text=self.chemin.split("/")[-1], anchor="w", command=lambda chemin=self.chemin: os.startfile(chemin))
        bouton_chemin.grid(row=1, column=2, sticky="ew")

        label_separateur = ttk.Label(frame, text="Séparateur")
        label_separateur.grid(row=2, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=2, column=1, sticky="ew", padx=10)

        self.barre_separateur = ttk.Entry(frame)
        self.barre_separateur.insert(0, ("Saut de ligne" if self.separateur == "SAUT_DE_LIGNE" else "Espace" if self.separateur == " " else self.separateur if self.separateur else ""))
        self.barre_separateur.bind("<KeyRelease>", lambda event: self.modifier_barre(self.barre_separateur))
        self.barre_separateur.grid(row=2, column=2, sticky="ew")

        if self.barre_separateur.get() == "":
            self.bouton_charger.config(state="disabled")

        label_colonne = ttk.Label(frame, text="Colonne")
        label_colonne.grid(row=3, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=3, column=1, sticky="ew", padx=10)

        df = pd.read_csv(self.chemin, sep=self.separateur, header=0 if self.var_en_tete.get() else None)
        self.menu_colonne = MenuDeroulant(frame, taille=5, items=df.columns, text=df.columns[self.colonne], command=self.detection_des_mandats)
        self.menu_colonne.selection_index = self.colonne
        self.menu_colonne.grid(row=3, column=2, sticky="ew")

        frame_en_tete = ttk.Frame(frame)

        frame_en_tete.grid_columnconfigure(0, weight=0)
        frame_en_tete.grid_columnconfigure(1, weight=1)
        label_en_tete = ttk.Label(frame_en_tete, text="Mon fichier comporte des entêtes :")
        label_en_tete.grid(row=0, column=0, sticky="w")

        checkbox_en_tete = ttk.Checkbutton(frame_en_tete, variable=self.var_en_tete)
        checkbox_en_tete.grid(row=0, column=1, sticky="w")
        checkbox_en_tete.bind("<Button-1>",  lambda event: self.modifier_checkbox(self.var_en_tete))

        frame_en_tete.grid(row=4, column=0, columnspan=3, sticky="ew")

        return frame
    
    def page_texte(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=0)
        frame.grid_rowconfigure(1, weight=1)
        frame.grid_rowconfigure(2, weight=1)
        frame.grid_columnconfigure(0, weight=0)
        frame.grid_columnconfigure(1, weight=0)
        frame.grid_columnconfigure(2, weight=1)

        self.label_nombre_mandats = ttk.Label(frame)
        self.label_nombre_mandats.grid(row=0, column=0, columnspan=3, padx=10, pady=(0, 10))

        thread = th.Thread(target=self.detection_des_mandats, daemon=True)
        thread.start()

        label_chemin = ttk.Label(frame, text="Fichier")
        label_chemin.grid(row=1, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=1, column=1, sticky="ew", padx=10)

        bouton_chemin = tk.Button(frame, text=self.chemin.split("/")[-1], anchor="w", command=lambda chemin=self.chemin: os.startfile(chemin))
        bouton_chemin.grid(row=1, column=2, sticky="ew")

        label_separateur = ttk.Label(frame, text="Séparateur")
        label_separateur.grid(row=2, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=2, column=1, sticky="ew", padx=10)

        self.barre_separateur = ttk.Entry(frame)
        self.barre_separateur.insert(0, ("Saut de ligne" if self.separateur == "SAUT_DE_LIGNE" else "Espace" if self.separateur == " " else self.separateur if self.separateur else ""))
        self.barre_separateur.bind("<KeyRelease>", lambda event: self.modifier_barre(self.barre_separateur))
        self.barre_separateur.grid(row=2, column=2, sticky="ew")

        if self.barre_separateur.get() == "":
            self.bouton_charger.config(state="disabled")

        return frame
    
    def pied_de_page(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)
        frame.grid_columnconfigure(2, weight=1)

        self.bouton_charger = ttk.Button(frame, text="Charger", command=self.charger)
        self.bouton_charger.grid(row=0, column=0, sticky="news", padx=(10, 5), pady=10)
        self.bouton_choisir_autre_fichier = ttk.Button(frame, text="Choisir un autre fichier", command=self.choisir_autre_fichier)
        self.bouton_choisir_autre_fichier.grid(row=0, column=1, sticky="news", padx=5, pady=10)
        self.bouton_annuler = ttk.Button(frame, text="Annuler", command=self.destroy)
        self.bouton_annuler.grid(row=0, column=2, sticky="news", padx=(5, 10), pady=10)

        return frame
    
    def detection_des_mandats(self):
        print(self.var_en_tete.get())

        try:
            separateur = self.barre_separateur.get()
            self.separateur = "SAUT_DE_LIGNE" if separateur.strip().lower() == "saut de ligne" else " " if separateur.strip().lower() == "espace" else separateur
        except:
            pass

        try:
            self.colonne = self.menu_colonne.get_index()
        except:
            pass

        self.label_nombre_mandats.config(text="Détection des mandats...")
        liste_candidats_mandats = utils.charger_liste_mandats(self.chemin, ("\n" if self.separateur == "SAUT_DE_LIGNE" else self.separateur), self.var_en_tete.get(), self.colonne)

        self.liste_mandats = []
        for candidat in liste_candidats_mandats:
            if candidat in self.racine.liste_mandats:
                self.liste_mandats.append(candidat)
        
        self.label_nombre_mandats.config(text=f"{len(self.liste_mandats)} mandat{"s" if len(self.liste_mandats) > 1 else ""} détecté{"s" if len(self.liste_mandats) > 1 else ""}.")

    def modifier_barre(self, barre):

        print(self.var_en_tete.get())

        if barre.get() != "":
            self.bouton_charger.config(state="normal")
        else:
            self.bouton_charger.config(state="disabled")

        if barre.get().strip().lower() == "saut de ligne":
            barre.delete(0, tk.END)
            barre.insert(0, "Saut de ligne")
        elif barre.get().strip().lower() == "espace":
            barre.delete(0, tk.END)
            barre.insert(0, "Espace")

        try:
            df = pd.read_csv(self.chemin, sep=barre.get(), header=0 if self.var_en_tete.get() else None)
            if self.colonne >= len(df.columns):
                self.colonne = 0
            self.menu_colonne.config(text=str(df.columns[self.colonne]))
            self.menu_colonne.set_items(df.columns)
            self.menu_colonne.hide_items()
        except:
            self.colonne = 0
            self.menu_colonne.config(text="")
            self.menu_colonne.set_items([])
            self.menu_colonne.hide_items()

        if len(barre.get()) > 0:
            self.detection_des_mandats()
        else:
            self.liste_mandats = []
            self.label_nombre_mandats.config(text="0 mandat détecté.")

    def modifier_checkbox(self, var_checkbox):

        self.var_en_tete.set(not self.var_en_tete.get())

        try:
            df = pd.read_csv(self.chemin, sep=self.barre_separateur.get(), header=0 if var_checkbox.get() else None)
            if self.colonne >= len(df.columns):
                self.colonne = 0
            self.menu_colonne.config(text=str(df.columns[self.colonne]))
            self.menu_colonne.set_items(df.columns)
            self.menu_colonne.hide_items()
        except:
            pass

        try:
            df = pd.read_excel(self.chemin, header=0 if var_checkbox.get() else None)
            if self.colonne >= len(df.columns):
                self.colonne = 0
            self.menu_colonne.config(text=str(df.columns[self.colonne]))
            self.menu_colonne.set_items(df.columns)
            self.menu_colonne.hide_items()
        except:
            pass
        

        self.detection_des_mandats()

        self.var_en_tete.set(not self.var_en_tete.get())

    def charger(self):
        self.detection_des_mandats()

        for mandat in self.liste_mandats:
            self.racine.ajouter_mandat_a_editer(mandat)

        self.mandats_ajoutes = self.liste_mandats
        self.destroy()

    def choisir_autre_fichier(self):
        self.destroy()
        self.racine.charger()


class PopUpArchiver(PopUp):
    def __init__(self, master, racine, on_finish, largeur, hauteur, paquet_standard, paquet_max):
        self.paquet_standard = paquet_standard
        self.paquet_max = paquet_max

        super().__init__(master, racine, on_finish, largeur, hauteur)

    def page(self):
        self.grid_rowconfigure(0, weight=0)
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=0)
        self.grid_columnconfigure(0, weight=1)

        frame_en_tete = self.en_tete()
        frame_en_tete.grid(row=0, column=0, sticky="news")
        frame_pied_de_page = self.pied_de_page()
        frame_pied_de_page.grid(row=2, column=0, sticky="news")
        self.frame_corps = self.corps()
        self.frame_corps.grid(row=1, column=0, sticky="new", padx=10, pady=10)

    def en_tete(self):
        frame = ttk.Frame(self)
        label_en_tete = ttk.Label(frame, text="Archiver les mandats", font=("Helvetica", 16))
        label_en_tete.pack(padx=10, pady=10)
        return frame
    
    def corps(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=0)
        frame.grid_rowconfigure(1, weight=1)
        frame.grid_rowconfigure(2, weight=1)
        frame.grid_columnconfigure(0, weight=0)
        frame.grid_columnconfigure(1, weight=0)
        frame.grid_columnconfigure(2, weight=1)

        self.label_nombre_mandats = ttk.Label(frame)
        self.label_nombre_mandats.grid(row=0, column=0, columnspan=3, padx=10, pady=(0, 10))

        thread = th.Thread(target=self.detection_des_mandats, daemon=True)
        thread.start()

        label_chemin = ttk.Label(frame, text="Fichier")
        label_chemin.grid(row=1, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=1, column=1, sticky="ew", padx=10)

        bouton_chemin = tk.Button(frame, text=self.chemin.split("/")[-1], anchor="w", command=lambda chemin=self.chemin: os.startfile(chemin))
        bouton_chemin.grid(row=1, column=2, sticky="ew")

        label_separateur = ttk.Label(frame, text="Séparateur")
        label_separateur.grid(row=2, column=0, sticky="w")
        ttk.Label(frame, text=":").grid(row=2, column=1, sticky="ew", padx=10)

        self.barre_separateur = ttk.Entry(frame)
        self.barre_separateur.insert(0, ("Saut de ligne" if self.separateur == "SAUT_DE_LIGNE" else "Espace" if self.separateur == " " else self.separateur if self.separateur else ""))
        self.barre_separateur.bind("<KeyRelease>", lambda event: self.modifier_barre(self.barre_separateur))
        self.barre_separateur.grid(row=2, column=2, sticky="ew")

        if self.barre_separateur.get() == "":
            self.bouton_charger.config(state="disabled")

        return frame
    
    def pied_de_page(self):
        frame = ttk.Frame(self)
        frame.grid_rowconfigure(0, weight=1)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(1, weight=1)

        self.bouton_charger = ttk.Button(frame, text="Archiver", command=self.charger)
        self.bouton_charger.grid(row=0, column=0, sticky="news", padx=(10, 5), pady=10)
        self.bouton_annuler = ttk.Button(frame, text="Annuler", command=self.destroy)
        self.bouton_annuler.grid(row=0, column=1, sticky="news", padx=(5, 10), pady=10)

        return frame
    



# LANCEMENT DE L'APPLICATION
if __name__ == "__main__":

    root = TkinterDnD.Tk()
    root.withdraw()

    def lancer_app(mandats, mandats_a_archiver, position):
        app = Application(root, mandats, mandats_a_archiver, position)

    splash = Demarrage(master=root, on_finish=lancer_app)
    root.mainloop()